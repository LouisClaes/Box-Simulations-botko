from datetime import datetime
from pathlib import Path

from docx import Document


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    now = datetime.now()
    datum = now.strftime("%d-%m-%Y")
    tijd = now.strftime("%H:%M")

    root = Path(__file__).resolve().parents[1]
    out_path = root / "Validatierapport_Botko_Closing_Logic.docx"

    doc = Document()
    doc.add_heading("Validatierapport Botko Overnight", level=0)
    add_par(doc, "Onderwerp: Verificatie marge-instellingen en automatische sluitlogica")
    add_par(doc, f"Datum: {datum}")
    add_par(doc, f"Tijd: {tijd}")

    add_heading(doc, "1. Doel", level=1)
    add_par(
        doc,
        "Dit document valideert of de overnight Botko-runner de juiste veiligheids- en "
        "operationele regels toepast, met focus op: (1) marge van 2 cm, en "
        "(2) automatische palletsluiting via lookahead-logica (huidige 4 + volgende 4 dozen).",
    )

    add_heading(doc, "2. Scope", level=1)
    add_bullets(
        doc,
        [
            "Runner: run_overnight_botko.py",
            "Sessie-orchestratie: simulator/session.py",
            "Marge-handhaving: simulator/bin_state.py",
            "Configuratie: config.py",
            "Functionele referentie: CLOSE_LOGIC_EXPLAINED.md",
        ],
    )

    add_heading(doc, "3. Samenvatting (Management)", level=1)
    add_par(
        doc,
        "Conclusie: de Botko overnight flow gebruikt de marge van 2 cm correct "
        "en bevat de automatische sluitlogica met lookahead van 4+4 dozen. "
        "De implementatie bevat daarnaast veiligheidsvoorwaarden: niet de laatste "
        "actieve pallet sluiten en alleen sluiten vanaf 50% vulling.",
    )

    add_heading(doc, "4. Resultaten", level=1)

    add_heading(doc, "4.1 Configuratie in Botko overnight", level=2)
    add_bullets(
        doc,
        [
            "BOTKO_SESSION_CONFIG gebruikt buffer_size = 8.",
            "BOTKO_SESSION_CONFIG gebruikt pick_window = 4.",
            "BOTKO_SESSION_CONFIG gebruikt close_policy = HeightClosePolicy(max_height=1800.0).",
            "BOTKO_PALLET is EUR-formaat 1200 x 800 x 2700 mm met 10 mm resolutie.",
            "De gebruikte bin-configuratie heeft margin = 20.0 mm (2 cm).",
        ],
    )
    add_par(
        doc,
        "Interpretatie: de overnight runner is geconfigureerd conform het 8-buffer / 4-grip model "
        "en gebruikt de palletparameters waar de margelogica op draait.",
    )

    add_heading(doc, "4.2 Margevalidatie (2 cm)", level=2)
    add_bullets(
        doc,
        [
            "Wall-margin wordt hard afgedwongen in BinState.get_height_at(...): "
            "bij schending van marge wordt een sentinelhoogte teruggegeven waardoor plaatsing faalt.",
            "Box-to-box marge wordt afgedwongen in BinState.is_margin_clear(...): "
            "z-overlappende footprint-uitbreiding met margin voorkomt te krappe plaatsingen.",
            "Configuratiebron voor marge is BinConfig.margin = 20.0 mm.",
        ],
    )
    add_par(
        doc,
        "Interpretatie: zowel tegen wanden als tussen dozen wordt de 2 cm-regel structureel toegepast.",
    )

    add_heading(doc, "4.3 Sluitlogica via lookahead (4+4)", level=2)
    add_bullets(
        doc,
        [
            "In _run_singlebin_step wordt bij 'geen plaatsing gevonden' eerst reject-counter verhoogd.",
            "Vervolgens wordt lookahead uitgevoerd: huidige pick-window (4) plus volgende 4 uit de buffer.",
            "Als geen van die 8 nog past: lookahead_stuck = True.",
            "Extra hard-stop: max_rejects_reached wanneer reject-counter >= 8.",
            "Sluiting gebeurt wanneer should_close = lookahead_stuck OR max_rejects_reached.",
            "Veiligheidsvoorwaarden: alleen sluiten bij meerdere actieve pallets en "
            "alleen pallets met >= 50% fill-rate.",
            "Bij sluiting wordt de volste kandidaatpallet gekozen, gesnapshot en vervangen door een verse pallet.",
        ],
    )
    add_par(
        doc,
        "Interpretatie: de gewenste lookahead-sluitlogica is aanwezig en operationeel, "
        "met aanvullende safeguards tegen ongewenste sluitingen.",
    )

    add_heading(doc, "4.4 Referentie-documentatie aanwezig", level=2)
    add_par(
        doc,
        "De logica is expliciet beschreven in CLOSE_LOGIC_EXPLAINED.md, inclusief de "
        "4+4-lookahead, reject-limiet en rationale.",
    )

    add_heading(doc, "5. Risico's en nuances", level=1)
    add_bullets(
        doc,
        [
            "Lookahead-sluiting treedt niet op in elk scenario: de veiligheidsfilters (>=50% fill, niet laatste pallet) zijn bewust strikt.",
            "De overnight runner gebruikt HeightClosePolicy als reguliere policy; de lookahead-logica is extra fallback in de sessieflow.",
        ],
    )

    add_heading(doc, "6. Eindoordeel", level=1)
    add_par(
        doc,
        "Goedgekeurd voor gebruik: de Botko overnight pipeline implementeert zowel de 2 cm margehandhaving "
        "als de automatische lookahead-sluitlogica conform de functionele beschrijving.",
    )

    add_heading(doc, "7. Bronnen (Code/Documentatie)", level=1)
    add_bullets(
        doc,
        [
            "python/CLOSE_LOGIC_EXPLAINED.md",
            "python/run_overnight_botko.py",
            "python/simulator/session.py",
            "python/simulator/bin_state.py",
            "python/config.py",
        ],
    )

    doc.save(out_path)
    print(f"Rapport aangemaakt: {out_path}")


if __name__ == "__main__":
    main()
